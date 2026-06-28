"""Regenerate the binary document fixtures (PDF/DOCX/XLSX/XLS).

These fixtures stand in for real clinic price lists published in office formats,
so the offline seed and the test-suite can exercise the PDF/DOCX/Excel parsers
deterministically without network access.

Run from the backend venv:

    python parsers/tests/fixtures/generate_documents.py

Authoring deps (NOT required at runtime): reportlab (PDF), xlwt (legacy .xls).
docx/xlsx are produced with python-docx / openpyxl which are already runtime deps.
"""
from __future__ import annotations

from pathlib import Path

HERE = Path(__file__).resolve().parent

# (service name as a clinic would write it, turnaround days or "", price KZT)
QARAGANDY_ROWS = [
    ("Общий анализ крови (ОАК)", "1", 2400),
    ("Биохимический анализ крови", "1", 8900),
    ("Глюкоза", "1", 1500),
    ("Гормон ТТГ", "1", 3200),
    ("С-реактивный белок (СРБ)", "1", 2800),
    ("Ферритин", "2", 3900),
    ("Витамин Д", "3", 9500),
    ("Коагулограмма", "2", 5600),
    ("Озонотерапия (1 процедура)", "", 6000),  # intentionally not in catalog
]

ATYRAU_ROWS = [
    ("Приём терапевта", 6500),
    ("Консультация кардиолога", 9000),
    ("Приём невролога", 8500),
    ("Приём гинеколога", 8000),
    ("УЗИ щитовидной железы", 7000),
    ("ЭКГ", 4000),
    ("Эхокардиография", 12000),
    ("Капельница", 5500),
    ("Иглорефлексотерапия (сеанс)", 7500),  # intentionally not in catalog
]

TARAZ_ROWS = [
    ("УЗИ органов брюшной полости", 9000),
    ("МРТ головного мозга", 28000),
    ("Рентген грудной клетки", 4500),
    ("Гастроскопия (ФГДС)", 15000),
    ("Колоноскопия", 22000),
    ("Маммография", 8000),
    ("Внутримышечная инъекция", 1200),
    ("Лечебный массаж", 6000),
    ("Холтер ЭКГ", 11000),
]


# Core PDF fonts (Helvetica) lack Cyrillic glyphs, so we must embed a TTF that
# has them. Try common system/bundled fonts; fail loudly if none is found.
_CYRILLIC_FONT_CANDIDATES = [
    "/System/Library/Fonts/Supplemental/Arial.ttf",
    "/Library/Fonts/Arial Unicode.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    "C:/Windows/Fonts/arial.ttf",
]


def _register_cyrillic_font() -> str:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    for path in _CYRILLIC_FONT_CANDIDATES:
        if Path(path).exists():
            pdfmetrics.registerFont(TTFont("Cyrillic", path))
            return "Cyrillic"
    raise RuntimeError(
        "No Cyrillic TTF font found. Install one (e.g. DejaVu Sans) or edit "
        "_CYRILLIC_FONT_CANDIDATES in this script."
    )


def make_pdf() -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font = _register_cyrillic_font()
    out = HERE / "qaragandy_clinic.pdf"
    doc = SimpleDocTemplate(str(out), pagesize=A4, title="Прайс-лист")
    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    title_style.fontName = font
    data = [["Наименование услуги", "Срок, дней", "Цена, тг"]]
    data += [[name, days, f"{price:,}".replace(",", " ")] for name, days, price in QARAGANDY_ROWS]
    table = Table(data, colWidths=[300, 80, 100])
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("FONTNAME", (0, 0), (-1, -1), font),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
            ]
        )
    )
    story = [
        Paragraph("Медцентр Сункар — Караганда. Прайс-лист 2026", title_style),
        Spacer(1, 12),
        table,
    ]
    doc.build(story)
    print("wrote", out.name)


def make_docx() -> None:
    import docx

    out = HERE / "atyrau_clinic.docx"
    document = docx.Document()
    document.add_heading("Клиника Жайык (Атырау) — прайс 2026", level=1)
    table = document.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Услуга"
    hdr[1].text = "Стоимость (тенге)"
    for name, price in ATYRAU_ROWS:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = str(price)
    document.save(str(out))
    print("wrote", out.name)


def make_xlsx() -> None:
    import openpyxl

    out = HERE / "taraz_clinic.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Прайс"
    # Title rows above the header exercise header auto-detection.
    ws.append(["Медцентр Аулие-Ата, г. Тараз"])
    ws.append(["Действует с 01.01.2026"])
    ws.append(["Наименование услуги", "Цена, тг"])
    for name, price in TARAZ_ROWS:
        ws.append([name, price])
    wb.save(str(out))
    print("wrote", out.name)


def make_xls() -> None:
    try:
        import xlwt
    except ImportError:
        print("skip sample.xls (xlwt not installed)")
        return
    out = HERE / "sample_clinic.xls"
    wb = xlwt.Workbook()
    ws = wb.add_sheet("Прайс")
    ws.write(0, 0, "Услуга")
    ws.write(0, 1, "Цена")
    rows = [("Общий анализ крови (ОАК)", 2300), ("Приём терапевта", 6000)]
    for r, (name, price) in enumerate(rows, start=1):
        ws.write(r, 0, name)
        ws.write(r, 1, price)
    wb.save(str(out))
    print("wrote", out.name)


if __name__ == "__main__":
    make_pdf()
    make_docx()
    make_xlsx()
    make_xls()
