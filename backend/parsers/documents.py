"""Tabular extraction from binary documents: PDF, DOCX, XLSX, XLS.

These helpers turn a document's bytes into a list of string rows
(``list[list[str]]``), which source parsers then map to :class:`Offer` rows.
They implement the ТЗ 3.1 requirement: "Поддержка форматов: HTML, PDF, DOCX,
Excel". HTML is handled separately by BeautifulSoup-based source parsers.

Each extractor degrades gracefully: a format whose optional dependency is
missing raises a clear ``DocumentError`` instead of crashing the pipeline, so
one broken format never stops ingestion of the others (fault-tolerance).
"""
from __future__ import annotations

import io
import logging
from collections.abc import Iterable

from parsers.utils import clean_text

logger = logging.getLogger(__name__)

#: Document formats this module can extract tabular rows from.
SUPPORTED_FORMATS = ("pdf", "docx", "xlsx", "xls")


class DocumentError(RuntimeError):
    """Raised when a document cannot be parsed (corrupt file or missing dep)."""


def _stringify_row(values: Iterable[object]) -> list[str]:
    """Normalize a raw row of mixed cell types into clean strings."""
    out: list[str] = []
    for v in values:
        if v is None:
            out.append("")
        elif isinstance(v, float) and v.is_integer():
            # openpyxl reads "5000" as 5000.0; keep it tidy for price parsing.
            out.append(str(int(v)))
        else:
            out.append(clean_text(str(v)))
    return out


def _drop_empty_rows(rows: list[list[str]]) -> list[list[str]]:
    return [r for r in rows if any(cell.strip() for cell in r)]


def extract_pdf_rows(data: bytes) -> list[list[str]]:
    """Extract rows from a PDF.

    Strategy: prefer ruled tables (``extract_tables``); if a page has none,
    fall back to splitting text lines on runs of 2+ spaces — common in price
    lists rendered without table borders.
    """
    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover - dependency declared in reqs
        raise DocumentError("pdfplumber is required to parse PDF documents") from exc

    rows: list[list[str]] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables() or []
            if tables:
                for table in tables:
                    for raw in table:
                        rows.append(_stringify_row(raw))
                continue
            # No ruled table on this page: parse text lines positionally.
            text = page.extract_text() or ""
            for line in text.splitlines():
                cells = [c for c in line.split("  ") if c.strip()]
                if len(cells) >= 2:
                    rows.append(_stringify_row(cells))
    return _drop_empty_rows(rows)


def extract_docx_rows(data: bytes) -> list[list[str]]:
    """Extract rows from a DOCX: every table row, plus 2+ column text lines."""
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise DocumentError("python-docx is required to parse DOCX documents") from exc

    document = docx.Document(io.BytesIO(data))
    rows: list[list[str]] = []
    for table in document.tables:
        for row in table.rows:
            rows.append(_stringify_row(cell.text for cell in row.cells))
    # Some clinics use tab-separated paragraphs instead of real tables.
    if not rows:
        for para in document.paragraphs:
            cells = [c for c in para.text.split("\t") if c.strip()]
            if len(cells) >= 2:
                rows.append(_stringify_row(cells))
    return _drop_empty_rows(rows)


def extract_xlsx_rows(data: bytes) -> list[list[str]]:
    """Extract rows from all sheets of a modern .xlsx workbook (openpyxl)."""
    try:
        import openpyxl
    except ImportError as exc:  # pragma: no cover
        raise DocumentError("openpyxl is required to parse XLSX documents") from exc

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    rows: list[list[str]] = []
    try:
        for ws in wb.worksheets:
            for raw in ws.iter_rows(values_only=True):
                rows.append(_stringify_row(raw))
    finally:
        wb.close()
    return _drop_empty_rows(rows)


def extract_xls_rows(data: bytes) -> list[list[str]]:
    """Extract rows from a legacy .xls workbook (xlrd)."""
    try:
        import xlrd
    except ImportError as exc:
        raise DocumentError(
            "xlrd is required to parse legacy .xls documents (pip install xlrd)"
        ) from exc

    book = xlrd.open_workbook(file_contents=data)
    rows: list[list[str]] = []
    for sheet in book.sheets():
        for r in range(sheet.nrows):
            rows.append(_stringify_row(sheet.row_values(r)))
    return _drop_empty_rows(rows)


_EXTRACTORS = {
    "pdf": extract_pdf_rows,
    "docx": extract_docx_rows,
    "xlsx": extract_xlsx_rows,
    "xls": extract_xls_rows,
}


def extract_rows(data: bytes, fmt: str) -> list[list[str]]:
    """Dispatch to the right extractor by format (``pdf``/``docx``/``xlsx``/``xls``)."""
    fmt = fmt.lower().lstrip(".")
    extractor = _EXTRACTORS.get(fmt)
    if extractor is None:
        raise DocumentError(f"Unsupported document format: {fmt!r}")
    return extractor(data)


def format_from_filename(name: str) -> str | None:
    """Return the document format implied by a filename extension, if supported."""
    suffix = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    return suffix if suffix in _EXTRACTORS else None
